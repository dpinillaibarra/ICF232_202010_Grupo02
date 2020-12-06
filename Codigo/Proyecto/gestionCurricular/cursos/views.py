import os
from django.shortcuts import render, redirect
from django.views.generic import (CreateView, DeleteView, DetailView, ListView,
                                  UpdateView, TemplateView)
from django.core.exceptions import ObjectDoesNotExist
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from datetime import date
from io import StringIO, BytesIO
from django.contrib.auth.mixins import LoginRequiredMixin
from .models import Curso, CursoFactory, MDocumento, MRubrica, Descripcion, Tarea, ActaFirmas
from .forms import DocumentoForma, RubricaForma, DescripcionForma, TareaForma, ActaForma
from django.http import HttpResponse, Http404
from django.db.models import Avg


class HomePageView(TemplateView):
    def get(self, request, **kwargs):
        return render(request, 'index.html', context=None)

class HomeCursosView(LoginRequiredMixin, TemplateView):
    def get(self, request, **kwargs):
        cursoFactory = CursoFactory()
        return render(request, 'cursos.html',{'cursos': cursoFactory.obtenerCursos()})
                        
class DetalleCursoView(LoginRequiredMixin, TemplateView):
    def get(self, request, **kwargs):
        cursoFactory = CursoFactory()
        sigla=kwargs["sigla"]
        return render(request, 'curso.html', {'curso': cursoFactory.getCurso(sigla)})

def Modelo_Form_Upload(request):
    if request.method == 'POST':
        form = DocumentoForma(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('/curso/ICF333/')
    else:
        form = DocumentoForma()
    return render(request, 'upload.html', {
        'form': form
    })

def Modelo_Form_Rubrica(request):
    if request.method == 'POST':
        form = RubricaForma(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('/curso/ICF333/')
    else:
        form = RubricaForma()
    return render(request, 'rubrica.html', {
        'form': form
    })


def Modelo_Form_Descripcion(request):
    form = DescripcionForma()
   
    if request.method == 'POST':
        form = DescripcionForma(request.POST)
        if form.is_valid():
            instancia = form.save(commit=False)
            instancia.save()
            return redirect('/curso/ICF333/')
    else:
        form = DescripcionForma()
    return render(request, 'descripcion.html', {
        'form': form
    })

def eliminar(request):
    eliminar_desc = Descripcion.objects.last()
    eliminar_desc.delete()
    return redirect(to="proyecto")

def EditarDescripcion(request):
    desc = Descripcion.objects.last()
    data = {
        'form':DescripcionForma(instance=desc)
    }
    if request.method == 'POST':
        desc_form = DescripcionForma(data=request.POST, instance=desc)
        if desc_form.is_valid():
            desc_form.save()
            data['form'] = desc_form
    return render(request, 'descripcion.html', data)

def nueva_descripcion(request):
    nueva_desc = Descripcion.objects.last()
    data = {
        'nueva_desc':nueva_desc
    }
    return render(request, 'proyecto.html', data)

def home_documentos(request):
    context = {
        'file': MDocumento.objects.all()
    }
    return render(request, 'documentos_subidos.html', context)

def eliminar_documento(request, id):
    documento = MDocumento.objects.get(id=id)
    documento.delete()

    return redirect(to="documentos_subidos")

def nueva_tarea(request):
    data = {
        'form':TareaForma()
    }
    if request.method == 'POST':
        formulario = TareaForma(request.POST)
        if formulario.is_valid():
            formulario.save()
            data['mensaje'] = "Tarea creada correctamente"

    return render(request, 'anadir_tarea.html', data)

def home_calificaciones(request):
    tareas = Tarea.objects.all()
    average = tareas.aggregate(Avg("nota"))["nota__avg"]

    if average is None:
        average = 0

    average = round(average, 2)

    print(average)

    data = {
        'tareas':tareas,
        'average':average
    }
    return render(request, 'calificaciones.html', data)

def modificar_tarea(request, id):
    tarea = Tarea.objects.get(id=id)
    data = {
        'form':TareaForma(instance=tarea)
    }
    if request.method == 'POST':
        formulario = TareaForma(data=request.POST, instance=tarea)
        if formulario.is_valid():
            formulario.save()
            data['mensaje'] = "Modificado correctamente"
            data['form'] = formulario
    return render(request, 'modificar_tarea.html', data)

def eliminar_tarea(request, id):
    tarea = Tarea.objects.get(id=id)
    tarea.delete()

    return redirect(to="calificaciones")

def home_calificaciones_alumno(request):
    tareas = Tarea.objects.all()
    average = tareas.aggregate(Avg("nota"))["nota__avg"]

    if average is None:
        average = 0

    average = round(average, 2)

    data = {
        'tareas':tareas,
        'average':average
    }
    return render(request, 'calificaciones_alumno.html', data )

def descargar(request, path):
    file_path = os.path.join(settings.MEDIA_ROOT, path)
    if os.path.exists(file_path):
        with open(file_path, 'rb')as fh:
            response = HttpResponse(fh.read(), context_type = "documento")
            response['Content-Disposition'] = 'inline;filename='+os.path.basename(file_path)
            return response
        
    raise Http404

def GenerarActa(request):
    data = {
        'form': ActaForma,
    }
    if request.method == 'POST':
        formulario = ActaForma(request.POST)
        if formulario.is_valid():
            formulario.save()
            data['mensaje'] = "Documento firmado correctamente"
    return render(request, 'acta_cierre.html', data)

def Acta(request):
    lista_profes= ActaFirmas.objects.all()
    data = {
        'profesores':lista_profes
    }

    return render(request, 'acta.html', data)

def TestDocumento(request):
    document = Document()

    document.add_picture((r'%s/static/Imagenes/foto1.png' % (settings.PROJECT_PATH)), width=Inches(4))
    document.add_paragraph()
    document.add_paragraph("%s" % date.today().strftime('%B, %d, %Y'))
    document.add_paragraph('Este es un texto de prueba')

    docx_title = "PrimerActa"
    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length

    return response

def cover_export(request):

    document = Document()
    profesores = ActaFirmas.objects.last()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Verdana'
    font.size = Pt(11)

    """ * FORMATO PARA PONER PALABRAS EN NEGRITA / ITALIC*
    r1 = document.add_paragraph()

    runner = r1.add_run("Este texto esta en negrita")
    runner.bold = True
    runner.italic = True
    """

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)


    document.add_picture('../static/unabfoto.png', width=Inches(1))
    l_p = document.paragraphs[-1]
    l_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    r1 = document.add_paragraph()
    r1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    runner = r1.add_run("Examen Final Proyecto de Título II")
    runner.bold = True
        
    r2 = document.add_paragraph()
    r2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    runner2 = r2.add_run("Ingeniería Civil lnformática")
    runner2.bold = True 

    document.add_paragraph()
    r3 = document.add_paragraph()
    r3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    runner3 = r3.add_run("ACTA APERTURA DEFENSA DE TÍTULO")
    runner3.bold = True
    runner3.underline = True
    runner3.center = True

    r4 = document.add_paragraph()
    r4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    runner4 = r4.add_run("Señores profesores, estudiantes y público en general")
    runner4.bold = True

    document.add_paragraph("En Santiago a 30 de Enero del año 2020, el alumno(a) "+ profesores.Rut_Alumno +" presenta ante esta")
    document.add_paragraph("comisión su examen final (Defensa) de Proyecto de Título II, correspondiente a la ")
    document.add_paragraph("exposición pública del proyecto desarrollado durante la asignatura.")
    document.add_paragraph("La comisión evaluadora se encuentra conformada por:")
    document.add_paragraph()
    document.add_paragraph("Profesor Guía:      "+ profesores.Profesor_guia +" ")
    #document.add_paragraph(profesores.Profesor_guia)
    document.add_paragraph("Profesor Invitado:  "+ profesores.Profesor_invitado_1 +"")
    #document.add_paragraph(profesores.Profesor_invitado_1)
    document.add_paragraph("Profesor Invitado:  "+ profesores.Profesor_invitado_2 +"")
    #document.add_paragraph(profesores.Profesor_invitado_2)

    document.add_paragraph()
    r5 = document.add_paragraph()
    runner5 = r5.add_run("El Alumno dispone de 20 minutos para realizar la defensa de su Proyecto y, ")
    runner5.italic = True

    r6 = document.add_paragraph()
    runner6 = r6.add_run("posteriormente, la comisión realizará una ronda de preguntas para calificar la ")
    runner6.italic = True

    r7 = document.add_paragraph()
    runner7 = r7.add_run("presentación.")
    runner7.italic = True

    r8 = document.add_paragraph()
    r8.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    runner8 = r8.add_run("ACTA CIERRE DEFENSA DE TÍTULO II")
    runner8.bold = True
    runner8.underline = True

    document.add_paragraph("En Santiago a 30 de Enero del año 2020, el alumno(a) "+ profesores.Rut_Alumno +" ha rendido el")
    document.add_paragraph("Examen de Proyecto de Título II.")
    document.add_paragraph()
    document.add_paragraph("La comisión ha determinado que el alumno(a) ha obtenido en la Defensa de su Proyecto")
    document.add_paragraph("de Título una evaluación de  7.0.")
    document.add_paragraph()
    document.add_paragraph(""+ profesores.Profesor_guia +", "+ profesores.upload_at.strftime('%Y-%m-%d %H:%M') +"     ""              " + profesores.Profesor_invitado_1 + ", " + profesores.upload_at.strftime('%Y-%m-%d %H:%M')  + ""     )
    document.add_paragraph("_________________________                ________________________________")
    document.add_paragraph("   Firma Profesor Guía			          Firma Profesor(es) invitado(s)")

    document.save('Acta_Cierre.docx')

#Guardamos el contenido y descargamos desde la pagina.

    document_data = BytesIO()
    document.save(document_data)
    document_data.seek(0)
    response = HttpResponse(
        document_data.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename = "Acta_Cierre.docx"'
    response['Content-Encoding'] = "UTF-8"

    return response
